"""Word (.docx) ファイルから特許テキストと図面を抽出するサービス。

【書類名】マーカーでセクションを分割し、以下のフィールドに格納する:
  abstract         ← 【書類名】要約書
  claims_text      ← 【書類名】特許請求の範囲
  description_text ← 【書類名】明細書
  biblio            ← 【書類名】特許願（なければ全文から抽出）
"""
import re
import logging
from pathlib import Path
from .pdf_importer import PatentDocument
from .jplatpat_scraper import parse_biblio

logger = logging.getLogger(__name__)

# 全文から【書類名】マーカーを検索して次のマーカーまでを1セクションとする
_SECTION_SPLIT_RE = re.compile(r"【書類名】([^\n【]+)", re.MULTILINE)


def _split_sections(full_text: str) -> dict[str, str]:
    """全文を【書類名】マーカーで分割してセクション辞書を返す。

    各セクションの内容は次の【書類名】の直前まで。
    """
    matches = list(_SECTION_SPLIT_RE.finditer(full_text))
    if not matches:
        return {}

    sections: dict[str, str] = {}
    for i, match in enumerate(matches):
        key = match.group(1).strip()
        content_start = match.end()
        content_end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        sections[key] = full_text[content_start:content_end].strip()

    return sections


async def import_word(path: Path | str) -> PatentDocument:
    """Word ファイルを読み込み、セクション別にテキスト・書誌情報・画像を返す。"""
    from docx import Document

    path = Path(path)
    logger.info(f"Word インポート開始: {path.name}")
    doc = Document(str(path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    images: list[bytes] = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                images.append(rel.target_part.blob)
            except Exception:
                pass

    sections = _split_sections(full_text)
    logger.info(f"検出セクション: {list(sections.keys())}")

    abstract = sections.get("要約書", "")
    claims_text = sections.get("特許請求の範囲", "")
    description_text = sections.get("明細書", "")

    # 【書類名】セクションが1つも見つからない場合は全文を description_text に格納
    if not (abstract or claims_text or description_text):
        logger.warning("【書類名】セクションが見つかりません。全文を詳細な説明として格納します。")
        description_text = full_text

    # 書誌情報は【特許願】セクションから抽出、なければ全文から
    biblio_source = sections.get("特許願", full_text)
    biblio = parse_biblio(biblio_source)

    # 【特許願】に発明の名称がない場合は【明細書】から補完
    if not biblio.get("title") and description_text:
        m = re.search(r"【発明の名称】\s*(.+)", description_text)
        if m:
            biblio["title"] = m.group(1).strip()

    # parse_biblio は【出願人】を探すが Word 書式は【特許出願人】のため別途補完
    if not biblio.get("applicant"):
        for source in [biblio_source, description_text]:
            m = re.search(r"【特許出願人】.*?【氏名又は名称】\s*([^\n]+)", source, re.DOTALL)
            if m:
                biblio["applicant"] = m.group(1).strip()
                break

    logger.info(f"書誌解析結果: title={biblio.get('title', '')!r}")

    return PatentDocument(
        text=full_text,
        images=images,
        metadata={"source_file": path.name},
        biblio=biblio,
        abstract=abstract,
        claims_text=claims_text,
        description_text=description_text,
    )
