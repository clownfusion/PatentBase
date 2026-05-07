"""Word / Excel ドキュメント生成サービス。

特許の分析結果を Word レポートおよび Excel サマリーとして出力する。
"""
from __future__ import annotations

import io
import json
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from backend.app.models.patent import Patent


# ─── Word レポート ──────────────────────────────────────────────────────────

def generate_word_report(patent: "Patent") -> bytes:
    """特許分析結果を Word (.docx) 形式で出力する。

    Returns:
        .docx ファイルのバイト列
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ページ余白
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def _add_heading(text: str, level: int) -> None:
        p = doc.add_heading(text, level=level)
        if p.runs:
            run = p.runs[0]
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
            elif level == 2:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
            else:
                run.font.size = Pt(11)

    def _add_label_value(label: str, value: str) -> None:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}：")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_value = p.add_run(value or "—")
        run_value.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    def _add_bullet(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # タイトルブロック
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("特許分析レポート")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    doc.add_paragraph()

    # 書誌情報
    _add_heading("書誌情報", 1)
    _add_label_value("特許番号", patent.patent_number or "—")
    _add_label_value("発明の名称", patent.title or "—")
    _add_label_value("出願人", patent.applicant or "—")
    _add_label_value("IPC 分類", patent.ipc_codes or "—")
    _add_label_value("データソース", patent.source or "—")

    # 要約
    if patent.abstract:
        _add_heading("要約（Abstract）", 2)
        p = doc.add_paragraph(patent.abstract)
        p.paragraph_format.space_after = Pt(6)

    # AI 分析結果
    if patent.analysis_status == "done":
        _add_heading("AI 分析結果", 1)

        if patent.summary:
            _add_heading("発明の概要", 2)
            p = doc.add_paragraph(patent.summary)
            p.paragraph_format.space_after = Pt(6)

        key_points = _parse_json_list(patent.key_points)
        if key_points:
            _add_heading("権利化ポイント", 2)
            for kp in key_points:
                _add_bullet(kp)

        claims = patent.claims_structured or []
        if claims:
            _add_heading("請求項の構造", 2)
            for claim in claims:
                num = claim.get("claim_number", "?")
                ctype = "独立請求項" if claim.get("claim_type") == "independent" else \
                        f"従属請求項（請求項{claim.get('depends_on', '?')}に従属）"
                summary_text = claim.get("summary", "")

                p = doc.add_paragraph()
                run_num = p.add_run(f"【請求項{num}】 ")
                run_num.bold = True
                run_num.font.size = Pt(10.5)
                run_type = p.add_run(f"（{ctype}）")
                run_type.font.size = Pt(9)
                run_type.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                p.paragraph_format.space_after = Pt(2)

                if summary_text:
                    p2 = doc.add_paragraph(f"  → {summary_text}")
                    p2.paragraph_format.space_after = Pt(2)
                    p2.paragraph_format.left_indent = Inches(0.2)
                    if p2.runs:
                        p2.runs[0].font.size = Pt(10)
                        p2.runs[0].italic = True

                components = claim.get("components", [])
                for comp in components:
                    comp_text = f"{comp.get('id', '')}.  {comp.get('description', '')}"
                    _add_bullet(comp_text)

                if claim.get("text"):
                    p3 = doc.add_paragraph(claim["text"])
                    p3.paragraph_format.left_indent = Inches(0.2)
                    p3.paragraph_format.space_after = Pt(4)
                    if p3.runs:
                        p3.runs[0].font.size = Pt(9.5)
                        p3.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        if patent.mermaid_diagram:
            _add_heading("構成図（Mermaid ソース）", 2)
            note = doc.add_paragraph(
                "以下のコードを https://mermaid.live または diagrams.net に貼り付けると図として表示できます。"
            )
            if note.runs:
                note.runs[0].font.size = Pt(9)
                note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            code_para = doc.add_paragraph(patent.mermaid_diagram)
            code_para.paragraph_format.left_indent = Inches(0.2)
            if code_para.runs:
                code_para.runs[0].font.name = "Courier New"
                code_para.runs[0].font.size = Pt(9)

    # 免責
    doc.add_paragraph()
    p = doc.add_paragraph(
        "※ 本レポートは PatentBase + Claude AI による自動生成です。"
        "実務利用の際は必ず専門家による確認を行ってください。"
    )
    if p.runs:
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Excel サマリー ─────────────────────────────────────────────────────────

def generate_excel_summary(patent: "Patent") -> bytes:
    """特許分析結果を Excel (.xlsx) 形式で出力する。

    シート構成:
      1. サマリー：書誌情報 + 発明概要 + 権利化ポイント
      2. 請求項：構造化請求項一覧
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()

    BLUE_DARK    = "1E40AF"
    BLUE_LIGHT   = "EFF6FF"
    BORDER_COLOR = "E2E8F0"

    thin = Side(style="thin", color=BORDER_COLOR)
    thin_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _header_row(ws, row: int, text: str) -> None:
        ws.row_dimensions[row].height = 22
        cell = ws.cell(row=row, column=1, value=text)
        cell.font = Font(name="游ゴシック", bold=True, color="FFFFFF", size=12)
        cell.fill = PatternFill("solid", fgColor=BLUE_DARK)
        cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)

    def _label_cell(ws, row: int, label: str) -> None:
        cell = ws.cell(row=row, column=1, value=label)
        cell.font = Font(name="游ゴシック", bold=True, size=10, color="475569")
        cell.fill = PatternFill("solid", fgColor="F1F5F9")
        cell.alignment = Alignment(horizontal="left", vertical="center", indent=1, wrap_text=True)
        cell.border = thin_border

    def _value_cell(ws, row: int, col: int, value, wrap: bool = False) -> None:
        cell = ws.cell(row=row, column=col, value=value)
        cell.font = Font(name="游明朝", size=10)
        cell.alignment = Alignment(horizontal="left", vertical="top", indent=1, wrap_text=wrap)
        cell.border = thin_border

    # ─── シート 1: サマリー
    ws1 = wb.active
    ws1.title = "サマリー"
    ws1.column_dimensions["A"].width = 18
    ws1.column_dimensions["B"].width = 36
    ws1.column_dimensions["C"].width = 36

    # タイトル行
    ws1.row_dimensions[1].height = 36
    tc = ws1.cell(row=1, column=1, value="特許分析レポート")
    tc.font = Font(name="游ゴシック", bold=True, size=16, color="FFFFFF")
    tc.fill = PatternFill("solid", fgColor="2563EB")
    tc.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws1.merge_cells("A1:C1")

    r = 3
    _header_row(ws1, r, "書誌情報"); r += 1
    for label, value in [
        ("特許番号",    patent.patent_number or "—"),
        ("発明の名称",  patent.title or "—"),
        ("出願人",      patent.applicant or "—"),
        ("IPC 分類",    patent.ipc_codes or "—"),
        ("データソース", patent.source or "—"),
    ]:
        ws1.row_dimensions[r].height = 18
        _label_cell(ws1, r, label)
        ws1.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
        _value_cell(ws1, r, 2, value, wrap=True)
        r += 1

    if patent.analysis_status == "done":
        r += 1
        _header_row(ws1, r, "発明の概要"); r += 1
        if patent.summary:
            ws1.row_dimensions[r].height = 60
            ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
            cell = ws1.cell(row=r, column=1, value=patent.summary)
            cell.font = Font(name="游明朝", size=10)
            cell.alignment = Alignment(wrap_text=True, vertical="top", indent=1)
            cell.border = thin_border
            r += 1

        r += 1
        _header_row(ws1, r, "権利化ポイント"); r += 1
        for i, kp in enumerate(_parse_json_list(patent.key_points), 1):
            ws1.row_dimensions[r].height = 30
            _label_cell(ws1, r, f"ポイント {i}")
            ws1.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            _value_cell(ws1, r, 2, kp, wrap=True)
            r += 1

    # ─── シート 2: 請求項
    claims = patent.claims_structured or []
    if claims:
        ws2 = wb.create_sheet(title="請求項")
        ws2.column_dimensions["A"].width = 10
        ws2.column_dimensions["B"].width = 14
        ws2.column_dimensions["C"].width = 12
        ws2.column_dimensions["D"].width = 36
        ws2.column_dimensions["E"].width = 50

        headers = ["請求項番号", "種別", "従属先", "技術的特徴（要旨）", "請求項全文（抜粋）"]
        ws2.row_dimensions[1].height = 24
        for col_i, lbl in enumerate(headers, 1):
            cell = ws2.cell(row=1, column=col_i, value=lbl)
            cell.font = Font(name="游ゴシック", bold=True, color="FFFFFF", size=10)
            cell.fill = PatternFill("solid", fgColor=BLUE_DARK)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        for row_i, claim in enumerate(claims, 2):
            ws2.row_dimensions[row_i].height = 40
            is_indep = claim.get("claim_type") == "independent"
            vals = [
                claim.get("claim_number", ""),
                "独立" if is_indep else "従属",
                str(claim.get("depends_on", "")) if claim.get("depends_on") else "",
                claim.get("summary", ""),
                (claim.get("text", "") or "")[:300],
            ]
            for col_i, val in enumerate(vals, 1):
                cell = ws2.cell(row=row_i, column=col_i, value=val)
                cell.font = Font(name="游明朝", size=9.5)
                cell.alignment = Alignment(wrap_text=True, vertical="top", indent=1)
                cell.border = thin_border
                if is_indep:
                    cell.fill = PatternFill("solid", fgColor=BLUE_LIGHT)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── 後方互換スタブ（旧 async インターフェース） ─────────────────────────────

async def export_analysis_to_word(analysis: dict, output_path) -> None:
    raise NotImplementedError("generate_word_report() を使用してください。")


async def export_analysis_to_excel(analysis: dict, output_path) -> None:
    raise NotImplementedError("generate_excel_summary() を使用してください。")


def get_drawio_xml(analysis: dict) -> str:
    return analysis.get("drawio_xml", "")


# ─── ユーティリティ ─────────────────────────────────────────────────────────

def _parse_json_list(value) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        return [str(v) for v in value]
    try:
        parsed = json.loads(value)
        if isinstance(parsed, list):
            return [str(v) for v in parsed]
    except (json.JSONDecodeError, TypeError):
        pass
    return [str(value)]
